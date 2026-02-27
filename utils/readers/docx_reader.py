#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
DOCX File Reader

This module provides functions to read text from .docx files.
Requires python-docx: pip install python-docx
"""

from docx import Document

def read_docx_plaintext(file_path):
    """
    读取 .docx 文件并返回纯文本内容
    """
    try:
        doc = Document(file_path)
        full_text = []
        
        # 读取所有段落
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)
        
        # 读取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text)
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        return "\n".join(full_text)
        
    except Exception as e:
        raise Exception(f"Error reading .docx file: {e}")

def read_docx_with_formatting(file_path):
    """
    读取 .docx 文件，保留基本格式信息（可选）
    """
    try:
        doc = Document(file_path)
        paragraphs_info = []
        
        for para in doc.paragraphs:
            para_info = {
                'text': para.text,
                'style': para.style.name if para.style else 'Normal',
                'runs': []
            }
            
            # 获取每个run的格式信息
            for run in para.runs:
                run_info = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline
                }
                if run_info['text']:
                    para_info['runs'].append(run_info)
            
            if para_info['text']:
                paragraphs_info.append(para_info)
        
        return paragraphs_info
        
    except Exception as e:
        raise Exception(f"Error reading .docx file: {e}")